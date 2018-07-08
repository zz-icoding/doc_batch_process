# !usr/bin/env python

'''
2018-6-22
B Day Special revision
show me the MAGIC!
增加功能：根据docx文件夹，生成文章汇总表

2018-6-19
rev 1.0
网络媒体编辑、汇总
筛选XXX(2016)年网络媒体文章，按指定顺序排序，生成一个word
'''

import os
import xlrd
import xlwt
import docx
import time
import datetime
import shutil
import string
import zipfile
import re
import warnings
import turtle

punc = string.punctuation + string.whitespace
digit = string.digits
refine_count = 0

#根据docx文件夹，生成文章汇总表
#docx文件位于original文件夹中
#当前文件夹中，需有：排序规则表.txt
def genDatabase(src_path = 'original', filename = 'default_database.xls'):
    #生成文章list
    origin_list = []
    date_pat = re.compile('\d{4}\D\d{1,2}\D\d{1,2}') # YYYYxMMxDD
    for doc_name in os.listdir(src_path):
        doc = docx.Document(os.path.join(src_path, doc_name))
        for para in doc.paragraphs[:: -1]: #搜索署名行
            if para.text.strip():
                sch = re.search(date_pat, para.text)
                if sch: #如匹配到日期
                    #日期
                    datestr = para.text[sch.start() : sch.end()]
                    split = []
                    for char in datestr:
                        if char not in digit:
                            split.append(char)
                    y, m, d = time.strptime(datestr,
                                            '%Y' + '%s' % split[0] + '%m' + '%s' % split[1] + '%d')[0:3]
                    date = datetime.datetime(y, m, d)
                    #媒体
                    sp = para.text.split(datestr)
                    media = sp[0].strip() if sp[0].strip() else sp[1].strip()
                    if not media: #如媒体字段为空
                        media = input('\n[%s]\nNO media name found. Please input a media name: '
                                      % doc_name).strip()
                        while not media:
                           media = input('Media is missing. Input again: ').strip()
                        warnings.warn('User generated [media:%s],\
                                      may NOT match [%s]' % (media, doc_name))
                else: #如未匹配到日期
                    if input('[%s]\nSigniture NOT found in the ending line. Discard this doc? Y/ N'
                             % doc_name).strip()[0].lower() == 'y':
                        break
                    else:
                        date = datetime.datetime(time.strptime(
                            input('Please input date (YYYY-MM-DD): ').strip(), '%Y-%m-%d')[0:3])
                        media = input('Please input media: ').strip()
                        warnings.warn('User generated [date:%s] and [media:%s],\
                                      may NOT match [%s]' % (date, media, doc_name))
                origin_list.append([media, date, doc_name])
                break
    #按文章编号排序
    for i in range(len(origin_list))[:0:-1]: #倒序搜索的重点为索引1，索引0不可再搜索：
        for j in range(len(origin_list))[i - 1::-1]: #因为内循环变成-1起始(相当于从end完整搜索)
            num_i = eval(origin_list[i][2].split('.')[0]) #每次内循环，num_i必须重算
            num_j = eval(origin_list[j][2].split('.')[0])
            if not isinstance(num_i, int):
                origin_list.append(origin_list.pop(i))
                break
            elif not isinstance(num_j, int):
                origin_list.append(origin_list.pop(j))
                continue
            if num_i < num_j:
                origin_list[i], origin_list[j] = origin_list[j], origin_list[i]
    #去除文章标题的序号和后缀名
    for item in origin_list: 
        item[2] = item[2].split('.')[1]
            
    #生成排序规则list
    seq_list = []
    seq_path = u'排序规则表.txt'
    while not os.path.exists(seq_path):
        input('Please copy [%s] to the root dir. When finished press Enter...')
    with open(u'排序规则表.txt') as f:
        for line in f:
            seq_list.append(line.strip())

    #生成文章汇总表
    genXls(origin_list, filename, seq_list)


#读取文章汇总表，并按媒体、日期排序
#媒体升序排列
#日期默认升序，可选降序
#文章汇总表：
#为.xlsx文件，或默认标题default_database.xls
#sheet 1，第2列媒体，第3列日期，第4列标题，从第2行开始
#sheet 2，第1列媒体排序规则，从第2行开始
def getDatabase(filename = r'网络媒体汇总.xlsx', date_reverse = False):
    if not os.path.exists(filename):
        for name in os.listdir():
            if name.endswith('.xlsx') or name == 'default_database.xls':
                filename = name
                break
        else:
            print('NO database found. Now generating...')
            filename = 'default_database.xls'
            genDatabase('original', filename)
    wb = xlrd.open_workbook(filename)

    #读取媒体文章汇总表
    w_sheet = wb.sheet_by_index(0)
    w_list = []
    for i in range(1, w_sheet.nrows):
        row = w_sheet.row_values(i, 1, 4)
        for i in range(len(row)):
            if isinstance(row[i], str):
                row[i] = row[i].strip().strip('《').strip('》')
        if isinstance(row[1], str): #处理日期列
            row[1] = row[1].split('（')[0].split('(')[0]
            split = []
            for char in row[1]:
                if char not in digit:
                    split.append(char)
            if len(split) == 2: # YYYYxMMxDD
                y, m, d = time.strptime(row[1],
                                        '%Y' + '%s' % split[0] + '%m' + '%s' % split[1] + '%d')[0:3]
            elif len(split) == 1: # YYYYxMM
                y, m, d = time.strptime(row[1], '%Y' + '%s' % split[0] + '%m')[0:3]
            else:
                raise Exception('date format error@ ', row)
            row[1] = datetime.datetime(y, m, d)
        else:
            row[1] = xlrd.xldate_as_datetime(row[1],0)
        w_list.append(row)

    #读取媒体名称排序规则表
    s_sheet = wb.sheet_by_index(1)
    s_list = []
    for i in range(1, s_sheet.nrows):
        s_list.append(s_sheet.cell_value(i, 0).strip().strip('《').strip('》'))

    #按媒体名称排序
    for s_item in s_list:
        for i in range(len(w_list))[::-1]: #倒序搜索，避免前序操作改变后续索引位置
            if w_list[i][0] == s_item:
                w_list.append(w_list.pop(i))
    for i in range(len(w_list))[::-1]: #将不包含在排序规则表中的媒体内容排到最后
        if w_list[i][0] not in s_list:
            w_list.append(w_list.pop(i))

    #按日期排序
    for i in range(len(w_list))[:0:-1]:
        for j in range(len(w_list))[i - 1::-1]:
            if w_list[j][0] == w_list[i][0]:
                if date_reverse: #降序
                    if w_list[j][1] < w_list[i][1]:
                        w_list[i], w_list[j] = w_list[j], w_list[i]                    
                else: #升序
                    if w_list[j][1] > w_list[i][1]:
                        w_list[i], w_list[j] = w_list[j], w_list[i]

    print('[%s] loaded' % filename)
    return w_list


#筛选XX(2016)年文章
def scrDocx(article_list, year = 2016, from_path = 'original', to_path = 'screened'):
    screen_list = []
    for item in article_list:
        if item[1].year == year:
            screen_list.append(item)
    genXls(screen_list, dst_path = 'screen_queued.xls')

    if os.path.exists(to_path):
        shutil.rmtree(to_path)
    os.mkdir(to_path)
    miss_count = 0
    #miss_list = []
    for i in range(len(screen_list)):
        miss = True
        mod_screen_name = screen_list[i][2] #去除特殊字符，便于比较
        for j in range(len(mod_screen_name))[::-1]: #倒序搜索
            if mod_screen_name[j] in punc:
                mod_screen_name = mod_screen_name[:j] + mod_screen_name[j + 1:] + ' '
                #末尾加空格，以确保索引正确
        mod_screen_name = mod_screen_name.strip()
        for dirpath, dirnames, filenames in os.walk(from_path):
            for filename in filenames:
                original_pathname = os.path.join(dirpath, filename)
                mod_filename = filename.rpartition('.')[0] #去除特殊字符，便于比较
                for j in range(len(mod_filename))[::-1]: #倒序搜索
                    if mod_filename[j] in punc:
                        mod_filename = mod_filename[:j] + mod_filename[j + 1:] + ' '
                mod_filename = mod_filename.strip()
                if mod_screen_name in mod_filename:
                    new_name = str(i + 1) + '.' + mod_screen_name + \
                               os.path.splitext(original_pathname)[1]
                    new_pathname = os.path.join(to_path, new_name)
                    silence = shutil.copy(original_pathname, new_pathname)
                    miss = False
        if miss: #通过miss来提取未找到的文件清单
            miss_count += 1
            #miss_list.append(screen_list[i])
            print('not found %2d: %d\t%s' % (miss_count, i + 1, screen_list[i]))

    print('\n %d docs copied to folder [%s]' % (len(screen_list) - miss_count, to_path))


#生成xls文件
#my_list写入Sheet1
#s_list(排序规则，如有)写入Sheet2
def genXls(my_list, dst_path = 'result.xls', s_list = None):
    wb = xlwt.Workbook()
    ws = wb.add_sheet('Sheet1')
    myFormat = xlwt.XFStyle()
    myFormat.num_format_str = 'yyyy/mm/dd'
    
    title = ['序号', '媒体', '日期', '标题']
    for i in range(len(title)):
        ws.write(0, i, title[i])

    rows = len(my_list)
    cols = len(my_list[0])
    for r in range(rows):
        ws.write(r + 1, 0, r + 1)
        for j in range(cols):
            ws.write(r + 1, j + 1, my_list[r][j], myFormat)

    if s_list:
        ws1 = wb.add_sheet('Sheet2')
        ws1.write(0, 0, '排序规则')
        for i in range(len(s_list)):
            ws1.write(i + 1, 0, s_list[i])
        
    wb.save(dst_path)
    print('[%s] generated' % dst_path)

    
#生成refined_doc文件
def refDocx(src_dir = 'screened', refine_folder = 'refined'):
    global refine_count
    if os.path.exists(refine_folder):
        shutil.rmtree(refine_folder)
    os.mkdir(refine_folder)

    #生成refined文件夹(单篇文档，格式修正)    
    temp_zip = 'temp.zip'
    temp_dir = 'temp'
    for doc_name in os.listdir(src_dir):
        old_doc = docx.Document(os.path.join(src_dir, doc_name))
        new_doc = docx.Document()
        setdocStyle(new_doc.styles['Normal']) #doc整体格式设置

        #写入refined_doc
        #每张图片作为单独段落
        fig_n = 0 #图片排序
        fig_unloaded = True #记录图片是否已经载入
        for para in old_doc.paragraphs:
            if para.text.strip():
                new_para = new_doc.add_paragraph()
            for run in para.runs:
                text = run.text.strip()
                if text: #copy非空文本到新文档
                    refined_text = refText(text, doc_name) #文本内容修订
                    new_run = new_para.add_run(refined_text)
                elif run.element.drawing_lst: #插入图片
                    #doc.inline_shapes无法识别非嵌入式图片，故直接从段落文字块判定有无图
                    if fig_unloaded: #如本doc还未载入图片，此处载入
                        silence = shutil.copy(os.path.join(src_dir, doc_name), temp_zip)
                        f = zipfile.ZipFile(temp_zip)
                        for file in f.namelist():
                            if file.startswith('word/media/image'):
                                silence = f.extract(file, temp_dir)
                        f.close()
                        img_dir = os.path.split(silence)[0]
                        img_list = os.listdir(img_dir)
                        fig_unloaded = False
                    for i in range(len(run.element.drawing_lst)):
                        fig_para = new_doc.add_paragraph() # 图片单独作为一段
                        fig_para.paragraph_format.first_line_indent = 0 #图片段落无缩进
                        fig_para.alignment = 1 #图片段落居中
                        fig_run = fig_para.add_run()
                        #if old_doc.inline_shapes: #如为内置图形（自动换行-嵌入式）
                        #    new_inlineshape = fig_run.add_picture(
                        #        os.path.join(img_dir, img_list[fig_n]),
                        #        old_doc.inline_shapes[fig_n].width,
                        #        old_doc.inline_shapes[fig_n].height)
                        #else: #如为其他图形
                        new_inlineshape = fig_run.add_picture(
                            os.path.join(img_dir, img_list[fig_n]),
                            width = docx.shared.Cm(13)) #默认图形宽度13cm
                        fig_n += 1
        if os.path.exists(temp_zip):
            os.remove(temp_zip)
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

        #refined doc格式调整
        new_doc.paragraphs[0].paragraph_format.first_line_indent = 0 #标题行无缩进
        new_doc.paragraphs[0].alignment = 1 #标题行居中
        new_doc.paragraphs[0].runs[0].font.bold = True #标题行加粗
        new_doc.paragraphs[1].insert_paragraph_before() #标题行后空行
        for i in range(len(new_doc.paragraphs))[::-1]:
            if new_doc.paragraphs[i].text.strip():
                new_doc.paragraphs[i].alignment = 2 #署名行右对齐(跳过末尾图片段落)
                break
        new_doc.add_paragraph() #文末空2行
        new_doc.add_paragraph() #文末空2行
        new_doc.save(os.path.join(refine_folder, doc_name))
    print('\nTotally %d refined.\n%d docs generated in [%s].' \
          % (refine_count, len(os.listdir(refine_folder)), refine_folder))


#生成汇总doc文件
def genDocx(src_dir = 'refined', dst_path = 'result.docx'):
    dst_doc = docx.Document()
    setdocStyle(dst_doc.styles['Normal']) #doc整体格式设置
    temp_zip = 'temp.zip'
    temp_dir = 'temp'
    for doc_name in os.listdir(src_dir):
        refined_doc = docx.Document(os.path.join(src_dir, doc_name))
        #doc如有图片，提取图片
        if refined_doc.inline_shapes:
            fig_n = 0 #图片排序
            silence = shutil.copy(os.path.join(src_dir, doc_name), temp_zip)
            f = zipfile.ZipFile(temp_zip)
            for file in f.namelist():
                if file.startswith('word/media/image'):
                    silence = f.extract(file, temp_dir)
            f.close()
            img_dir = os.path.split(silence)[0]
            img_list = os.listdir(img_dir)
        #标题段落
        title_para = dst_doc.add_paragraph(refined_doc.paragraphs[0].text)
        title_para.paragraph_format.first_line_indent = 0 #标题行无缩进
        title_para.alignment = 1 #标题行居中
        title_para.runs[0].font.bold = True #标题行加粗
        #其他段落
        for para in refined_doc.paragraphs[1 :]:
            other_para = dst_doc.add_paragraph()
            if para.runs:
                for run in para.runs:
                    new_run = other_para.add_run(run.text)
                    if run.element.drawing_lst: #如是图片，插入图片段
                        for i in range(len(run.element.drawing_lst)):
                            other_para.paragraph_format.first_line_indent = 0 #图片段落无缩进
                            other_para.alignment = 1 #图片段落居中
                            new_inlineshape = new_run.add_picture(
                                os.path.join(img_dir, img_list[fig_n]),
                                refined_doc.inline_shapes[fig_n].width,
                                refined_doc.inline_shapes[fig_n].height)
                            fig_n += 1
        #署名段
        for i in range(len(dst_doc.paragraphs))[::-1]:
            if dst_doc.paragraphs[i].text.strip():
                dst_doc.paragraphs[i].alignment = 2 #署名行右对齐(跳过末尾图片、空行)
                break

        if os.path.exists(temp_zip):
            os.remove(temp_zip)
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)

    dst_doc.save(dst_path)
    print('\nSummarized doc [%s] generated.' % dst_path)

   
# 文本内容完善
# 修正英文引号
# 纠正前引号、回引号使用
# 纠正英文逗号，句号
def refText(text, pathname, output = False):
    global refine_count
    temp_text = text
    # 引号修正
    quote_revised = False
    quote_en = ["'", '"']
    quote_cn_s = ['‘', '“'] #前引号
    quote_cn_e = ['’', '”'] #回引号
    for j in range(len(quote_en)):
        quote_s = True #单引号、双引号分别前引号搜索起点
        for i in range(len(temp_text)):
            if temp_text[i] == quote_en[j]:
                quote_revised = True
                refine_count += 1
                if quote_s:
                    temp_text = temp_text[:i] + quote_cn_s[j] + temp_text[i + 1 :]
                    quote_s = False
                else:
                    temp_text = temp_text[:i] + quote_cn_e[j] + temp_text[i + 1 :]
                    quote_s = True
    if quote_revised and output:
        print('\n[%s]\nOriginal:\t%s' % (pathname, text))
        print('quote revised:\t%s' % temp_text)

    #逗号、句号修正
    comma_period_revised = False
    cp_en = [',', '.']
    cp_cn = ['，', '。']
    for j in range(len(cp_en)):
        for i in range(len(temp_text)):
            idx_s = []
            idx_e = []
            new = []
            if temp_text[i] == cp_en[j]:
                if i < (len(temp_text) - 1) and (temp_text[i - 1] in digit) \
                   and (temp_text[i + 1] in digit):
                    pass # 剔除小数点, 千分位
                elif temp_text[i - 1] in digit:
                    pass # 剔除序号项
                elif temp_text[i - 1] == '.': 
                    pass # 配合下节使用，跳过正确的多点线（除第一个点外）
                elif i < (len(temp_text) - 1) and temp_text[i + 1] == '.': # 检查多点线
                    end = 0
                    for k in range(i + 1, len(temp_text)):
                        if temp_text[k] != '.':
                            end = k - 1
                            break
                    print('\n%s (%s) %s'
                          % (temp_text[:i], temp_text[i : end + 1], temp_text[end + 1:]))
                    if input('is (%s) right? Y/N ' % temp_text[i : end + 1]).lower() == 'y':
                        pass
                    else:
                        new.insert(0, input('Please input the right content to replace (%s): '
                                    % temp_text[i : end + 1])) #倒序，避免索引错误
                        idx_s.insert(0, i) #倒序，避免索引错误
                        idx_e.insert(0, end + 1) #倒序，避免索引错误
                        refine_count += 1
                else:
                    refine_count += 1
                    comma_period_revised = True
                    temp_text = temp_text[:i] + cp_cn[j] + temp_text[i + 1 :]
    if idx_s:
        for i in range(len(idx_s)):
            temp_text = temp_text[: idx_s[i]] + new[i] + temp_text[idx_e[i]:]
    if comma_period_revised and output:
        if not quote_revised:
            print('\n[%s]\nOriginal:\t%s' % (pathname, text))
        print('comma / period revised:\t%s' % temp_text)
    
    return temp_text


#doc总体格式设置
def setdocStyle(doc_style_obj,
             font_name = 'Hans', #字体名称
             font_size = docx.shared.Pt(10.5), #字体大小
             first_line_indent_n = 2, #首行缩进2字
             space_before = 0, #段前间距0
             space_after = 0, #段后间距0
             line_spacing = 1.0): #行距单倍         
    doc_style_obj.font.name = 'Times New Roman' 
    doc_style_obj.font.size = docx.shared.Pt(10.5) 
    doc_style_obj.paragraph_format.first_line_indent = \
    doc_style_obj.font.size * first_line_indent_n
    doc_style_obj.paragraph_format.space_before = space_before
    doc_style_obj.paragraph_format.space_after = space_after
    doc_style_obj.paragraph_format.line_spacing = line_spacing

#生日快乐
def love():
    def func(x, y):
        main()
    turtle.title('喵喵子程序')
    lv=turtle.Turtle()
    lv.hideturtle()
    lv.getscreen().bgcolor('light blue')
    lv.color('yellow','red')
    lv.pensize(1)
    lv.speed(1)
    lv.up()
    lv.goto(0,-150)
    #开始画爱心
    lv.down()
    lv.begin_fill()
    lv.goto(0, -150)
    lv.goto(-175.12, -8.59)
    lv.left(140)
    pos = []
    for i in range(19):
        lv.right(10)
        lv.forward(20)
        pos.append((-lv.pos()[0], lv.pos()[1]))
    for item in pos[::-1]:
        lv.goto(item)
    lv.goto(175.12, -8.59)
    lv.goto(0, -150)
    lv.left(50)
    lv.end_fill()
    #写字
    lv.up()
    lv.goto(0, 80)
    lv.down()
    lv.write("喵喵",font=(u"方正舒体",36,"normal"),align="center")
    lv.up()
    lv.goto(0, 0)
    lv.down()
    lv.write("生日快乐！",font=(u"方正舒体",48,"normal"),align="center")
    lv.up()
    lv.goto(100, -210)
    lv.down()
    lv.write("点我见证奇迹",font=(u"华文琥珀",26,"bold"),align="right")
    lv.up()
    lv.goto(160, -190)
    lv.resizemode('user')
    lv.shapesize(4, 4, 10)
    lv.color('red', 'red')
    lv.onclick(func)
    lv.showturtle()


def main():
    w_list = getDatabase(filename = r'网络媒体汇总.xlsx', date_reverse = True)
    genXls(w_list, dst_path = 'all_queued.xls')
    scrDocx(w_list, year = 2016, from_path = 'original', to_path = 'screened')
    refDocx(src_dir = 'screened', refine_folder = 'refined')
    genDocx(src_dir = 'refined', dst_path = 'result.docx')

if __name__ == '__main__':
    if datetime.date.today() == datetime.date(2018, 6, 22):
        love()
    else:
        main()
